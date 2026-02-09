
import os
from docx import Document
import pandas as pd

# Create Word Document
doc_path = "data/docs/general/Guide_v1.docx"
doc = Document()
doc.add_heading('General Policy Guide', 0)
doc.add_paragraph('This is a sample Word document for testing conversion.')
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph('We use python-docx to extract this text.')
doc.save(doc_path)
print(f"Created {doc_path}")

# Create Excel File
excel_path = "data/docs/finance/Stats_2025.xlsx"
df = pd.DataFrame({
    'Quarter': ['Q1', 'Q2', 'Q3', 'Q4'],
    'Revenue': [100, 150, 200, 250],
    'Cost': [80, 90, 110, 130]
})
df.to_excel(excel_path, index=False)
print(f"Created {excel_path}")
