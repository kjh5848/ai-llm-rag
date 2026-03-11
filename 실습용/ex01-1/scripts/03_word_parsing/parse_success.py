"""
[실습 3 - Word 스타일 추출]
워드 문서의 자체 스타일(Heading)을 인식하여 마크다운 계층 구조(#, ##)로 매핑합니다.
"""
import os
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(BASE_DIR, '../../data/docs/sec/SEC_보안규정_v1.0.docx')
OUTPUT_FILE = os.path.join(BASE_DIR, '../../parsed_data/success_results/03_SEC_보안규정_success.md')

def convert_docx_to_md(docx_path, md_path):
    print(f"\n📝 Word 스타일 변환 시작: {os.path.basename(docx_path)}")
    os.makedirs(os.path.dirname(md_path), exist_ok=True)
    
    doc = Document(docx_path)
    md_content = []
    
    # 1. 단락(Paragraph) 텍스트와 스타일 매핑
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        style_name = para.style.name
        
        # 제목(Heading) 스타일 변환
        if style_name == 'Title':
            md_content.append(f"# {text}")
        elif style_name.startswith('Heading'):
            level = style_name.split()[-1]
            try:
                md_content.append(f"{'#' * (int(level)+1)} {text}")
            except ValueError:
                md_content.append(f"## {text}")
        # 리스트(Bullet) 스타일 변환
        elif style_name == 'List Bullet':
            md_content.append(f"- {text}")
        else:
            md_content.append(text)
            
    md_content.append("\n## [내부 점검표 (Table)]\n")
            
    # 2. 표(Table) 추출
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            md_content.append("| " + " | ".join(row_data) + " |")
            if i == 0: # 헤더 구분선 추가
                md_content.append("| " + " | ".join(["---"] * len(row_data)) + " |")
        md_content.append("\n")

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_content))
    print(f"✅ 마크다운 변환 완료: {os.path.basename(md_path)}")

if __name__ == "__main__":
    convert_docx_to_md(INPUT_FILE, OUTPUT_FILE)
