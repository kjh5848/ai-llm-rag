import os
import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Frame, PageTemplate, Image, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from PIL import Image as PilImage, ImageDraw, ImageFont

def html_to_pdf_playwright(html_path, pdf_path):
    import time
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("Playwright is not installed. Skipping PDF generation.")
        return
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            url = f"file://{os.path.abspath(html_path)}"
            page.goto(url)
            time.sleep(1)
            page.pdf(path=pdf_path, format="A4", print_background=True, margin={"top": "0cm", "bottom": "0cm", "left": "0cm", "right": "0cm"})
            browser.close()
    except Exception as e:
        print(f"Playwright PDF Error: {e}")


# Set Base Path
BASE_DIR = os.path.join(os.path.dirname(__file__), '../data/docs')

# Register Korean Font for reportlab
font_path = os.path.join(BASE_DIR, 'NanumGothic.ttf')
if not os.path.exists(font_path):
    print("한글 폰트(NanumGothic.ttf) 다운로드 중...")
    import urllib.request
    url = "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf"
    urllib.request.urlretrieve(url, font_path)

pdfmetrics.registerFont(TTFont('Korean', font_path))

# Create custom styles
styles = getSampleStyleSheet()
title_style = ParagraphStyle(
    name='KoreanTitle', fontName='Korean', fontSize=24, leading=28,
    alignment=1, spaceAfter=30, textColor=colors.HexColor("#1e3a8a")
)
heading_style = ParagraphStyle(
    name='KoreanHeading', fontName='Korean', fontSize=14, leading=18,
    spaceAfter=15, textColor=colors.HexColor("#0f172a"),
    borderPadding=(5, 5, 5, 5), backColor=colors.HexColor("#f1f5f9")
)
body_style = ParagraphStyle(
    name='KoreanBody', fontName='Korean', fontSize=10, leading=16,
    spaceAfter=12, textColor=colors.HexColor("#334155")
)
small_style = ParagraphStyle(
    name='KoreanSmall', fontName='Korean', fontSize=8, leading=10,
    spaceAfter=5, textColor=colors.HexColor("#64748b")
)


def create_hr_rules_pdf():
    html_path = os.path.join(BASE_DIR, 'hr', 'HR_취업규칙_v1.0.html')
    pdf_path = os.path.join(BASE_DIR, 'hr', 'HR_취업규칙_v1.0.pdf')
    
    template_path = '/Users/nomadlab/Desktop/김주혁/workspace/coding-study/ai-llm-rag/실습용/ex01/docs/html/HR_사내규정_다단.html'
    with open(template_path, 'r', encoding='utf-8') as tf:
        html_content = tf.read()
    
    html_content = html_content.replace('사내 인사 규정 (발췌)', '취업규칙 (다단 편집형)')
    html_content = html_content.replace('Metacoding Inc.', '메타코딩')
    
    # 텍스트 추출기가 레이아웃을 구분할 수 있게 힌트를 주는 테두리 선(가로선, 세로선)을 모두 제거
    html_content = html_content.replace('border-bottom: 2px solid #333;', '')
    html_content = html_content.replace('column-rule: 1px solid #ddd;', '')
    html_content = html_content.replace('border-bottom: 1px solid #eee;', '')
    html_content = html_content.replace('border-top: 1px solid #eee;', '')
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    html_to_pdf_playwright(html_path, pdf_path)
    print(f"생성 완료: {pdf_path}")
def create_fin_sales_excel():
    path = os.path.join(BASE_DIR, 'finance', 'FIN_2025_상반기_매출현황.xlsx')
    
    import pandas as pd
    import numpy as np
    from openpyxl.styles import PatternFill, Font, Border, Side, Alignment

    arrays = [
        ['Global'] * 12,
        ['아태지역(APAC)'] * 4 + ['미주(NA)'] * 4 + ['유럽(EMEA)'] * 4,
        ['한국(서울)', '한국(부산)', '일본(도쿄)', '대만(타이베이)', '미국(뉴욕)', '미국(텍사스)', '캐나다(밴쿠버)', '멕시코(시티)', '영국(런던)', '프랑스(파리)', '독일(베를린)', '이탈리아(로마)']
    ]
    index = pd.MultiIndex.from_tuples(list(zip(*arrays)), names=['Business Unit', 'Region', 'Branch'])
    
    columns_arrays = [['1분기(Q1)'] * 3 + ['2분기(Q2)'] * 3, ['1월', '2월', '3월', '4월', '5월', '6월']]
    columns = pd.MultiIndex.from_tuples(list(zip(*columns_arrays)), names=['Quarter', 'Month'])

    df = pd.DataFrame(np.random.randint(50000, 999999, (12, 6)), index=index, columns=columns)
    
    writer = pd.ExcelWriter(path, engine='openpyxl')
    df.to_excel(writer, sheet_name='Sales_Data')
    worksheet = writer.sheets['Sales_Data']
    
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    header_fill = PatternFill("solid", fgColor="1F497D")
    header_font = Font(bold=True, color="FFFFFF")
    idx_fill = PatternFill("solid", fgColor="DCE6F1")
    idx_font = Font(bold=True, color="1F497D")
    
    for row in worksheet.iter_rows():
        for cell in row:
            if type(cell).__name__ != 'MergedCell':
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center')
                if cell.row <= 3 and cell.column > 3:
                    cell.fill = header_fill
                    cell.font = header_font
                elif cell.row > 3 and cell.column <= 3:
                    cell.fill = idx_fill
                    cell.font = idx_font
                elif cell.row > 3 and cell.column > 3:
                    cell.number_format = '#,##0'
                    cell.alignment = Alignment(horizontal='right', vertical='center')
    
    idx_names = ["Business Unit", "Region", "Branch"]
    for i, name in enumerate(idx_names, 1):
        cell = worksheet.cell(row=3, column=i, value=name)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for col, width in zip('ABCDEFGHI', [15, 20, 20, 15, 15, 15, 15, 15, 15]):
        worksheet.column_dimensions[col].width = width

    writer.close()
    print(f"생성 완료: {path}")
def create_fin_budget_excel():
    # 4. 결재란이 포함된 문서형 엑셀 (양식 포맷 강화)
    path = os.path.join(BASE_DIR, 'finance', 'FIN_부서별_예산기안서.xlsx')
    wb = Workbook()
    ws = wb.active
    ws.title = "예산기안서_최종"

    thin_border = Border(left=Side(style='thin'), 
                         right=Side(style='thin'), 
                         top=Side(style='thin'), 
                         bottom=Side(style='thin'))

    # 대제목
    ws.merge_cells("A1:H2")
    cell = ws["A1"]
    cell.value = "2025년도 주요사업부 예산 집행 기안서"
    cell.font = Font(size=20, bold=True, color="000080")
    cell.alignment = Alignment(horizontal="center", vertical="center")

    # 상단 메타데이터 (RAG 파싱 방해 요소)
    ws["A4"] = "문서번호:"
    ws["B4"] = "FIN-BDG-2025-001"
    ws["A5"] = "기안부서:"
    ws["B5"] = "재무전략실"
    ws["A6"] = "작성일자:"
    ws["B6"] = "2025. 02. 20"
    
    for row in range(4, 7):
        ws.cell(row=row, column=1).font = Font(bold=True)
        ws.cell(row=row, column=1).alignment = Alignment(horizontal="right")

    # 우측 결재란 만들기
    ws.merge_cells("F4:F6")
    ws["F4"] = "결\n\n재"
    ws["F4"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws["F4"].border = thin_border

    headers_approv = ["기안", "검토", "승인"]
    for col_idx, h in enumerate(headers_approv, 7):
        cell = ws.cell(row=4, column=col_idx, value=h)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border
        cell.fill = PatternFill("solid", fgColor="EFEFEF")
        
        ws.merge_cells(start_row=5, start_column=col_idx, end_row=6, end_column=col_idx)
        blank_cell = ws.cell(row=5, column=col_idx)
        blank_cell.border = thin_border

    ws.row_dimensions[5].height = 20
    ws.row_dimensions[6].height = 20

    # 본문 헤더
    header_row = 9
    headers = ["#", "본부명", "부서명", "예산 계정", "Q1 배정액(원)", "Q2 배정액(원)", "상반기 총액", "집행 사유 (비고)"]
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col_idx, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="4F81BD")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border

    # 본문 데이터 세트
    data = [
        [1, "경영지원본부", "인사총무팀", "복리후생비", 25000000, 30000000, "=E10+F10", "전사 춘계 워크샵 비용 포함"],
        [2, "경영지원본부", "재무전략실", "지급수수료", 12000000, 15000000, "=E11+F11", "외부 회계감사 수수료"],
        [3, "기술개발본부", "백엔드개발팀", "소프트웨어구입", 85000000, 45000000, "=E12+F12", "클라우드 서비스 및 데이터베이스 연간 라이선스"],
        [4, "기술개발본부", "AI리서치팀", "연구개발비", 120000000, 180000000, "=E13+F13", "고성능 GPU 서버 증설 및 API 토큰 비용"],
        [5, "글로벌영업본부", "북미영업팀", "여비교통비", 45000000, 60000000, "=E14+F14", "CES 등 해외 전시회 참가 항공/숙박비"],
        [6, "글로벌영업본부", "마케팅팀", "광고선전비", 200000000, 350000000, "=E15+F15", "신규 프로덕트 런칭 글로벌 디지털 캠페인"]
    ]
    
    for r_idx, row_data in enumerate(data, header_row + 1):
        for c_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.border = thin_border
            if c_idx in [5, 6, 7]: # 숫자 포맷
                cell.number_format = '#,##0'
                cell.alignment = Alignment(horizontal="right")
            elif c_idx == 8: # 비고란 왼쪽 정렬
                 cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

    # 합계
    sum_row = header_row + len(data) + 1
    ws.merge_cells(start_row=sum_row, start_column=1, end_row=sum_row, end_column=4)
    cell = ws.cell(row=sum_row, column=1, value="총 합계")
    cell.font = Font(bold=True)
    cell.fill = PatternFill("solid", fgColor="D9D9D9")
    cell.alignment = Alignment(horizontal="center")
    cell.border = thin_border
    
    for col in range(2, 5):
        ws.cell(row=sum_row, column=col).border = thin_border

    for col in [5, 6, 7]:
        cell = ws.cell(row=sum_row, column=col, value=f"=SUM({ws.cell(row=header_row+1, column=col).coordinate}:{ws.cell(row=sum_row-1, column=col).coordinate})")
        cell.font = Font(bold=True)
        cell.border = thin_border
        cell.number_format = '#,##0'
    
    ws.cell(row=sum_row, column=8).border = thin_border
    ws.cell(row=sum_row, column=8).fill = PatternFill("solid", fgColor="D9D9D9")

    # 컬럼 너비
    col_widths = {'A': 5, 'B': 15, 'C': 15, 'D': 15, 'E': 15, 'F': 15, 'G': 15, 'H': 35}
    for col, width in col_widths.items():
        ws.column_dimensions[col].width = width

    wb.save(path)
    print(f"생성 완료: {path}")


def create_sec_rules_docx():
    # 5. 스타일 및 특수기호 워드 문서 (내용 및 서식 강화)
    path = os.path.join(BASE_DIR, 'security', 'SEC_보안규정_v1.0.docx')
    doc = Document()
    
    # Title
    t = doc.add_heading('전사 정보 보안 규정 및 가이드라인', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('문서 버젼: v1.0.3\n최종 배포일: 2025-02-20\n보안 등급: 회사 극비(Top Secret)').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Heading 1
    doc.add_heading('제 1 장. 망분리 및 접근 통제', level=1)
    
    # Heading 2
    doc.add_heading('1.1. 논리적 망분리 환경 지침', level=2)
    p = doc.add_paragraph('회사의 핵심 개발 서버 및 DB는 인터넷이 차단된 ')
    p.add_run('CDE(Cardholder Data Environment)').bold = True
    p.add_run(' 영역에 배치되며, 이 영역에 접근하기 위해서는 반드시 사내 인가된 VPN을 거쳐야 한다.')
    
    doc.add_heading('1.2. 비밀번호 생성 규칙 (강제)', level=2)
    doc.add_paragraph('시스템 접근용 비밀번호는 다음의 복잡도 요구사항을 모두 충족해야 한다.')
    
    # List Bullet
    doc.add_paragraph('☑ 영문 거대문자(A-Z) 및 소문자(a-z) 혼용 필수', style='List Bullet')
    doc.add_paragraph('☑ 숫자(0-9) 및 특수기호(!@#$%^&*) 1개 이상 반드시 포함', style='List Bullet')
    doc.add_paragraph('☐ 최소 12자리 이상 길이 (권장 16자리)', style='List Bullet')
    doc.add_paragraph('☒ (절대 금지 사항) 생년월일, 폰번호 등 개인정보 기반 유추 가능한 문자열 금지', style='List Bullet')

    doc.add_heading('제 2 장. 부서별 보안 점검 항목 (중첩 표 예시)', level=1)
    doc.add_paragraph('AI 모델이 표 내부의 특수 기호와 줄바꿈을 제대로 파싱(Parsing)하는지 실험하기 위한 목적으로 설계된 [별도 첨부] 점검표이다.')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1' # Word built-in style
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '부서 권한'
    hdr_cells[1].text = '점검 주기'
    hdr_cells[2].text = '상세 보안 점검 지표'
    hdr_cells[3].text = '담당자 확인(서명)'

    # Row 1
    row_cells = table.add_row().cells
    row_cells[0].text = '일반 임직원\n(Level 1)'
    row_cells[1].text = '월 1회\n(매월 말일)'
    
    # Cell 내부에 리스트 텍스트 작성 (중첩 구조 흉내)
    idx_text = (
        "1. 백신 프로그램 정의 파일 최신 여부\n"
        "2. 사내 메신저 자동 로그인 해제 여부\n"
        "3. 화면 보호기(10분) 비밀번호 잠금 설정"
    )
    row_cells[2].text = idx_text
    row_cells[3].text = '[   ] 양호\n[   ] 미흡'

    # Row 2
    row_cells = table.add_row().cells
    row_cells[0].text = 'IT 개발본부\n(Level 3)'
    row_cells[1].text = '주 1회\n(매주 금요일)'
    idx_text2 = (
        "• Git 저장소 SSH 키 암호화 보관 확인\n"
        "• AWS IAM 토큰 다중 인증(MFA) 활성화\n"
        "• 개발/상용망 간 불법 포트포워딩 검사"
    )
    row_cells[2].text = idx_text2
    row_cells[3].text = '[   ] 양호\n[   ] 미흡'

    doc.save(path)
    print(f"생성 완료: {path}")


def create_ops_strategy_pdf():
    # 6. 복합 프레젠테이션 (가로형 PPT 스타일 - LLM Vision/Refinement 용)
    path = os.path.join(BASE_DIR, 'ops', 'OPS_신규서비스_런칭전략.pdf')
    doc = SimpleDocTemplate(path, pagesize=landscape(A4), rightMargin=20, leftMargin=20, topMargin=30, bottomMargin=30)
    
    elements = []
    
    title_ppt = ParagraphStyle(
        name='TitlePPT', fontName='Korean', fontSize=32, leading=36,
        alignment=1, spaceAfter=20, textColor=colors.HexColor("#A855F7")
    )
    
    ## ----- PAGE 1: 아키텍처 슬라이드 ----- ##
    elements.append(Paragraph("<b>2026 초거대 AI RAG 시스템 엔터프라이즈 런칭 전략</b>", title_ppt))
    elements.append(Paragraph("혁신을 넘어서, 기업 내 모든 문서를 지식화하는 완벽한 파이프라인의 구축", ParagraphStyle(name='sub', fontName='Korean', fontSize=16, alignment=1, textColor=colors.grey, spaceAfter=20)))

    data = [
        [
         Paragraph("<font color='white'><b>Phase 1: 데이터 수집 🌊</b></font><br/><br/><font color='#e2e8f0'>- 사내 흩어진 PDF, Word 모음<br/>- ERP, Notion API 연동<br/>- OCR 및 Vision 기술 적용</font>", ParagraphStyle(name='b1', fontName='Korean', fontSize=13, leading=20, textColor=colors.white)),
         Paragraph("<font size=32 color='#94a3b8'>➔</font>", ParagraphStyle(name='arr', alignment=1)),
         Paragraph("<font color='white'><b>Phase 2: RAG 파이프라인 🧠</b></font><br/><br/><font color='#e2e8f0'>- Markdown 구조 복원 변환<br/>- Semantic Chunking 분할<br/>- BGE-M3 텍스트 임베딩 모델</font>", ParagraphStyle(name='b2', fontName='Korean', fontSize=13, leading=20, textColor=colors.white)),
         Paragraph("<font size=32 color='#94a3b8'>➔</font>", ParagraphStyle(name='arr', alignment=1)),
         Paragraph("<font color='white'><b>Phase 3: 에이전트 서비스 🚀</b></font><br/><br/><font color='#e2e8f0'>- ChromaDB 벡터 유사도 검색<br/>- DeepSeek-R1 추론 모델<br/>- 기업용 Slack/Teams 챗봇 배포</font>", ParagraphStyle(name='b3', fontName='Korean', fontSize=13, leading=20, textColor=colors.white))
        ]
    ]
    
    t = Table(data, colWidths=[8.2*cm, 1.5*cm, 8.2*cm, 1.5*cm, 8.2*cm], style=[
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BACKGROUND', (0,0), (0,0), colors.HexColor("#2563eb")), # Blue
        ('BACKGROUND', (2,0), (2,0), colors.HexColor("#7c3aed")), # Purple
        ('BACKGROUND', (4,0), (4,0), colors.HexColor("#059669")), # Green
        ('ROUNDEDCORNERS', [10, 10, 10, 10]), 
        ('PADDING', (0,0), (-1,-1), 15),
        ('BOTTOMPADDING', (0,0), (-1,-1), 15),
        ('TOPPADDING', (0,0), (-1,-1), 15),
    ])
    elements.append(t)
    elements.append(Spacer(1, 10))
    
    elements.append(Paragraph("<b>⚠️ 아키텍처 구현 시 주의사항 (Critical Constraints)</b>", heading_style))
    elements.append(Paragraph("위 그림과 같이 화살표와 텍스트 박스로 이루어진 프레젠테이션(PPT) 문서를 단순한 PDF 텍스트 추출기(예: PyMuPDF, pdfplumber)로 밀어 넣으면, 좌우의 텍스트가 마구잡이로 배열되어 <b>'데이터 수집 사내 흩어진 PDF RAG 파이프라인 Markdown...'</b> 처럼 알아볼 수 없는 쓰레기 데이터(Garbage)가 됩니다.<br/>이러한 문서는 반드시 <b>Vision-Language Model (VLM)</b> 에 이미지를 통째로 밀어 넣어 마크다운으로 구조를 복원(Refinement)하라고 명령해야만 검색 가능한 지식으로 탈바꿈할 수 있습니다.", ParagraphStyle(name='body2', fontName='Korean', fontSize=12, leading=18)))
    
    elements.append(Spacer(1, 25))
    elements.append(Paragraph("Page 1 / 2 | 메타코딩 Confidential", small_style))

    from reportlab.platypus.doctemplate import PageBreak
    elements.append(PageBreak())

    ## ----- PAGE 2: 로드맵 및 예산 슬라이드 ----- ##
    elements.append(Paragraph("<b>2026 하반기 통합 로드맵 및 소요 예산 (Roadmap & Budget)</b>", title_ppt))
    
    # 간트 차트를 모방한 복잡한 테이블 생성
    elements.append(Paragraph("<b>1. Phase별 마일스톤 (Milestones)</b>", heading_style))
    
    roadmap_data = [
        ['태스크 (Task)', 'Q1 (1~3월)', 'Q2 (4~6월)', 'Q3 (7~9월)', 'Q4 (10~12월)', '담당 조직'],
        ['1. 전사 데이터 수집 파이프라인 연동', '진행중 (In-Progress)', '', '', '', '데이터엔지니어링팀'],
        ['2. 문서 파싱용 Vision Model 튜닝', '', '예정 (Planned)', '', '', 'AI Research팀'],
        ['3. 벡터 DB (Chroma) 인프라 구축', '', '예정 (Planned)', '예정 (Planned)', '', '인프라보안팀'],
        ['4. 사내 챗봇 (Slack/Teams) UI 연동', '', '', '', '예정 (Planned)', '백엔드개발팀 / UX팀']
    ]
    
    t_road = Table(roadmap_data, colWidths=[6.8*cm, 4*cm, 4*cm, 4*cm, 4*cm, 4.4*cm], style=[
        ('FONTNAME', (0,0), (-1,-1), 'Korean'),
        ('FONTSIZE', (0,0), (-1,-1), 11),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0f172a")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        # 간트차트 색상 블록 시뮬레이션
        ('BACKGROUND', (1, 1), (1, 1), colors.HexColor("#3b82f6")), # Q1 Blue
        ('TEXTCOLOR', (1, 1), (1, 1), colors.white),
        ('BACKGROUND', (2, 2), (2, 2), colors.HexColor("#f59e0b")), # Q2 Orange
        ('BACKGROUND', (2, 3), (3, 3), colors.HexColor("#10b981")), # Q2-Q3 Green
        ('BACKGROUND', (4, 4), (4, 4), colors.HexColor("#8b5cf6")), # Q4 Purple
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
    ])
    elements.append(t_road)
    elements.append(Spacer(1, 5))

    elements.append(Paragraph("<b>2. 하반기 소요 예산 추정치 (단위: 백만원)</b>", heading_style))
    
    budget_data = [
        ['구분', '항목', '예산액', '비고 (산출 근거)'],
        ['인건비', 'AI 모델러 / 백엔드 엔지니어 충원', '350', '시니어 2명, 주니어 1명 신규 채용'],
        ['인프라비', 'GPU 서버 증설 (A100 x 4)', '200', '자체 LLM 및 Vision 모델 서빙용'],
        ['라이선스', 'VLM API (GPT-4V, DeepSeek 등)', '150', '월간 1억 토큰 기준 산정'],
        ['합계', '', '700', '']
    ]
    t_budget = Table(budget_data, colWidths=[3.5*cm, 8.5*cm, 3.5*cm, 11.7*cm], style=[
        ('FONTNAME', (0,0), (-1,-1), 'Korean'),
        ('FONTSIZE', (0,0), (-1,-1), 11),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#475569")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('SPAN', (0, 4), (1, 4)), # 합계 셀 병합
        ('BACKGROUND', (0, 4), (-1, 4), colors.HexColor("#e2e8f0")),
        ('FONTNAME', (0, 4), (-1, 4), 'Korean'), 
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
    ])
    elements.append(t_budget)
    elements.append(Spacer(1, 5))
    
    elements.append(Paragraph("✔ 본 슬라이드 또한 OCR/VLM 성능 평가를 위한 다중 표(Multi-table) 혼합 문서 구조를 가지고 있습니다. 표 내부에 배경색이 채워져 있거나 빈 셀이 있는 경우 단순 텍스트 파서에서는 Null 값 혹은 표가 깨지는 현상이 발생합니다.", ParagraphStyle(name='body3', fontName='Korean', fontSize=12, leading=18)))

    elements.append(Spacer(1, 5))
    elements.append(Paragraph("Page 2 / 2 | 메타코딩 Confidential", small_style))

    doc.build(elements)
    print(f"생성 완료: {path}")


def create_ops_image():
    # 7. 화려한 매출 차트 이미지 (Vision/VLM 적용 난이도 향상)
    path = os.path.join(BASE_DIR, 'ops', 'OPS_매출현황_v1.0.png')
    width, height = 800, 500
    img = PilImage.new('RGB', (width, height), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    
    try:
        font_title = ImageFont.truetype(font_path, 32)
        font_label = ImageFont.truetype(font_path, 18)
        font_val = ImageFont.truetype(font_path, 14)
    except:
        font_title = ImageFont.load_default()
        font_label = ImageFont.load_default()
        font_val = ImageFont.load_default()
    
    # Title & Subtitle
    d.text((40, 30), "2024~2025 핵심 플랫폼 MAU 및 매출 성장 추이", font=font_title, fill=(30, 41, 59))
    d.text((40, 75), "* 단위: MAU(만 명) / 매출액(십억 원)", font=font_label, fill=(100, 116, 139))
    
    # 그리드 선 그리기
    for y in range(150, 451, 75):
        d.line([(80, y), (700, y)], fill=(226, 232, 240), width=1)
        d.text((40, y-10), f"{int((450-y)*1.5)}", font=font_label, fill=(148, 163, 184))

    # 데이터
    quarters = ['24.Q1', '24.Q2', '24.Q3', '24.Q4', '25.Q1(E)', '25.Q2(E)']
    mau_data = [120, 180, 240, 300, 350, 420]
    rev_data = [80, 110, 165, 210, 270, 330]
    
    bar_width = 30
    spacing = 80
    start_x = 120
    base_y = 450
    
    for i, (q, mau, rev) in enumerate(zip(quarters, mau_data, rev_data)):
        x = start_x + (i * spacing)
        
        # MAU 막대 (파란색)
        h_mau = mau * 0.7
        d.rectangle([x, base_y - h_mau, x + bar_width, base_y], fill=(59, 130, 246))
        d.text((x+2, base_y - h_mau - 20), str(mau), font=font_val, fill=(30, 64, 175))
        
        # 매출 막대 (보라색)
        h_rev = rev * 0.7
        d.rectangle([x + bar_width, base_y - h_rev, x + bar_width*2, base_y], fill=(168, 85, 247))
        d.text((x + bar_width+2, base_y - h_rev - 20), str(rev), font=font_val, fill=(107, 33, 168))
        
        # X축 라벨
        d.text((x + 10, base_y + 15), q, font=font_label, fill=(71, 85, 105))

    # 꺾은선 (성장률 가상 선)
    points = [(start_x + bar_width + i*spacing, base_y - (m+r)*0.35 - 50) for i, (m, r) in enumerate(zip(mau_data, rev_data))]
    d.line(points, fill=(239, 68, 68), width=3)
    for px, py in points:
         d.ellipse([px-4, py-4, px+4, py+4], fill=(239, 68, 68))

    # 범례
    d.rectangle([600, 80, 620, 100], fill=(59, 130, 246))
    d.text((630, 82), "MAU", font=font_label, fill=(0,0,0))
    d.rectangle([600, 110, 620, 130], fill=(168, 85, 247))
    d.text((630, 112), "매출액", font=font_label, fill=(0,0,0))
    d.line([(595, 150), (625, 150)], fill=(239, 68, 68), width=3)
    d.text((630, 142), "성장률", font=font_label, fill=(0,0,0))

    img.save(path)
    print(f"생성 완료: {path}")

def create_hr_security_scan_pdf():
    html_path = os.path.join(BASE_DIR, 'hr', 'HR_정보보안서약서.html')
    pdf_path = os.path.join(BASE_DIR, 'hr', 'HR_정보보안서약서.pdf')
    
    template_path = '/Users/nomadlab/Desktop/김주혁/workspace/coding-study/ai-llm-rag/실습용/ex01/docs/html/SEC_보안규정_스캔.html'
    with open(template_path, 'r', encoding='utf-8') as tf:
        html_content = tf.read()
    
    html_content = html_content.replace('보안 규정 지침서', '정보보안 서약서 (스캔본)')
    html_content = html_content.replace('2026-SEC-001', '2026-HR-SEC-002')
    html_content = html_content.replace('Metacoding Inc.', '메타코딩')
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    html_to_pdf_playwright(html_path, pdf_path)
    print(f"생성 완료: {pdf_path}")

if __name__ == "__main__":
    print("--- 실무형 모의 문서 생성기 시작 ---")
    create_hr_rules_pdf()
    create_hr_security_scan_pdf()
    create_fin_sales_excel()
    create_fin_budget_excel()
    create_sec_rules_docx()
    create_ops_strategy_pdf()
    create_ops_image()
    print("--- 생성 완료 ---")
